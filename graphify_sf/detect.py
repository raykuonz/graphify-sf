import fnmatch
import hashlib
import json
import os
import sys
from enum import Enum
from pathlib import Path

try:
    import pathspec  # gitwildmatch — correct .gitignore/.forceignore semantics
except ImportError:  # pragma: no cover - pathspec is a core dependency
    pathspec = None


class SFFileType(str, Enum):
    APEX = "apex"
    TRIGGER = "trigger"
    FLOW = "flow"
    OBJECT = "object"
    FIELD = "field"
    CHILD_OBJECT = "child_object"
    LAYOUT = "layout"
    VISUALFORCE = "visualforce"
    LWC_BUNDLE = "lwc_bundle"
    AURA_BUNDLE = "aura_bundle"
    PROFILE = "profile"
    PERMISSIONSET = "permissionset"
    CONFIG = "config"
    AUTOMATION = "automation"
    EXPERIENCE = "experience"
    AGENTFORCE = "agentforce"


# Compound extensions MUST be checked first
_COMPOUND_EXT_MAP = {
    ".flow-meta.xml": SFFileType.FLOW,
    ".object-meta.xml": SFFileType.OBJECT,
    ".field-meta.xml": SFFileType.FIELD,
    ".validationRule-meta.xml": SFFileType.CHILD_OBJECT,
    ".recordType-meta.xml": SFFileType.CHILD_OBJECT,
    ".listView-meta.xml": SFFileType.CHILD_OBJECT,
    ".compactLayout-meta.xml": SFFileType.CHILD_OBJECT,
    ".businessProcess-meta.xml": SFFileType.CHILD_OBJECT,  # E2: target for RecordType refs
    ".layout-meta.xml": SFFileType.LAYOUT,
    ".profile-meta.xml": SFFileType.PROFILE,
    ".permissionset-meta.xml": SFFileType.PERMISSIONSET,
    ".permissionsetgroup-meta.xml": SFFileType.PERMISSIONSET,
    ".labels-meta.xml": SFFileType.CONFIG,
    ".md-meta.xml": SFFileType.CONFIG,
    ".settings-meta.xml": SFFileType.CONFIG,
    ".namedCredential-meta.xml": SFFileType.CONFIG,
    ".externalService-meta.xml": SFFileType.CONFIG,
    ".connectedApp-meta.xml": SFFileType.CONFIG,
    ".app-meta.xml": SFFileType.CONFIG,
    ".tab-meta.xml": SFFileType.CONFIG,
    ".flexipage-meta.xml": SFFileType.CONFIG,
    ".testSuite-meta.xml": SFFileType.CONFIG,
    ".remoteSite-meta.xml": SFFileType.CONFIG,
    ".remoteSiteSetting-meta.xml": SFFileType.CONFIG,  # A3: correct SFDX extension
    ".externalDataSource-meta.xml": SFFileType.CONFIG,  # A4
    ".authprovider-meta.xml": SFFileType.CONFIG,  # A7
    ".cspTrustedSite-meta.xml": SFFileType.CONFIG,  # A7
    ".corsWhitelistOrigins-meta.xml": SFFileType.CONFIG,  # A7
    ".sharingRules-meta.xml": SFFileType.CONFIG,  # B2
    ".sharingSet-meta.xml": SFFileType.CONFIG,  # B2
    ".restrictionRule-meta.xml": SFFileType.CONFIG,  # B5
    ".duplicateRule-meta.xml": SFFileType.CONFIG,  # B5
    ".matchingRule-meta.xml": SFFileType.CONFIG,  # B5
    ".fieldSet-meta.xml": SFFileType.CHILD_OBJECT,  # E1: FieldSet
    ".globalValueSet-meta.xml": SFFileType.CONFIG,  # E1: GlobalValueSet
    ".resource-meta.xml": SFFileType.CONFIG,  # D3: StaticResource
    ".quickAction-meta.xml": SFFileType.CONFIG,  # D3: QuickAction
    ".role-meta.xml": SFFileType.CONFIG,
    ".site-meta.xml": SFFileType.EXPERIENCE,
    ".network-meta.xml": SFFileType.EXPERIENCE,
    ".workflow-meta.xml": SFFileType.AUTOMATION,
    ".approvalProcess-meta.xml": SFFileType.AUTOMATION,
    ".escalationRules-meta.xml": SFFileType.AUTOMATION,
    ".assignmentRules-meta.xml": SFFileType.AUTOMATION,
    ".autoResponseRules-meta.xml": SFFileType.AUTOMATION,
    # Agentforce
    ".bot-meta.xml": SFFileType.AGENTFORCE,
    ".botVersion-meta.xml": SFFileType.AGENTFORCE,
    ".genAiPlugin-meta.xml": SFFileType.AGENTFORCE,
    ".genAiFunction-meta.xml": SFFileType.AGENTFORCE,
    ".genAiPlannerBundle-meta.xml": SFFileType.AGENTFORCE,
    ".aiAuthoringBundle-meta.xml": SFFileType.AGENTFORCE,
    ".promptTemplate-meta.xml": SFFileType.AGENTFORCE,
}

_SIMPLE_EXT_MAP = {
    ".cls": SFFileType.APEX,
    ".trigger": SFFileType.TRIGGER,
    ".page": SFFileType.VISUALFORCE,
    ".component": SFFileType.VISUALFORCE,
}


class DocFileType(str, Enum):
    DOCUMENT = "document"
    PAPER = "paper"
    IMAGE = "image"


_DOC_EXTENSIONS = {
    ".md": DocFileType.DOCUMENT,
    ".mdx": DocFileType.DOCUMENT,
    ".txt": DocFileType.DOCUMENT,
    ".rst": DocFileType.DOCUMENT,
    ".html": DocFileType.DOCUMENT,
    ".htm": DocFileType.DOCUMENT,
}
_PAPER_EXTENSIONS = {
    ".pdf": DocFileType.PAPER,
}
_OFFICE_EXTENSIONS = {
    ".docx": DocFileType.DOCUMENT,
    ".xlsx": DocFileType.DOCUMENT,
}
_IMAGE_EXTENSIONS = {
    ".png": DocFileType.IMAGE,
    ".jpg": DocFileType.IMAGE,
    ".jpeg": DocFileType.IMAGE,
    ".gif": DocFileType.IMAGE,
    ".webp": DocFileType.IMAGE,
    ".svg": DocFileType.IMAGE,
}

_SKIP_DIRS = {
    ".sfdx",
    "node_modules",
    "__pycache__",
    ".git",
    "graphify-sf-out",
    # Agentic-tooling scaffolding and repo-meta dirs — never contain real SF metadata.
    # Sample .cls files and skill templates here otherwise create phantom nodes.
    ".agents",
    ".claude",
    ".cursor",
    ".github",
    ".husky",
    ".vscode",
    ".omc",
    ".venv",
    ".ruff_cache",
    ".pytest_cache",
    "dist",
    "build",
}


def _compound_suffix(path: Path) -> str:
    """Return the longest matching compound suffix, or the simple suffix if none match."""
    suffixes = path.suffixes
    for i in range(len(suffixes), 0, -1):
        compound = "".join(suffixes[-i:])
        if compound in _COMPOUND_EXT_MAP:
            return compound
    return path.suffix if path.suffix else ""


def _classify_file(path: Path) -> SFFileType | None:
    """Classify a single file by compound extension."""
    suffix = _compound_suffix(path)
    if suffix in _COMPOUND_EXT_MAP:
        return _COMPOUND_EXT_MAP[suffix]
    if suffix in _SIMPLE_EXT_MAP:
        return _SIMPLE_EXT_MAP[suffix]
    return None


def _is_lwc_bundle(bundle_dir: Path) -> bool:
    """Return True if this directory is an LWC bundle."""
    if not bundle_dir.is_dir():
        return False
    name = bundle_dir.name
    js_file = bundle_dir / f"{name}.js"
    return js_file.exists()


def _is_aura_bundle(bundle_dir: Path) -> bool:
    """Return True if this directory is an Aura bundle."""
    if not bundle_dir.is_dir():
        return False
    name = bundle_dir.name
    cmp_file = bundle_dir / f"{name}.cmp"
    return cmp_file.exists()


def _load_sfgraphignore(root: Path) -> list[str]:
    """Load .graphifysfignore patterns from the project root."""
    ignore_file = root / ".graphifysfignore"
    if not ignore_file.exists():
        return []
    patterns = []
    for line in ignore_file.read_text(encoding="utf-8", errors="ignore").splitlines():
        line = line.strip()
        if line and not line.startswith("#"):
            patterns.append(line)
    return patterns


def _is_ignored(path: Path, root: Path, patterns: list[str]) -> bool:
    """Return True if the path matches any .graphifysfignore pattern (fnmatch)."""
    if not patterns:
        return False
    try:
        rel = str(path.relative_to(root))
    except ValueError:
        return False
    for pattern in patterns:
        if fnmatch.fnmatch(rel, pattern):
            return True
        if fnmatch.fnmatch(path.name, pattern):
            return True
    return False


# ---------------------------------------------------------------------------
# .gitignore / .forceignore support (gitwildmatch semantics via pathspec)
# ---------------------------------------------------------------------------
# Files honored by default. A user can opt out entirely with respect_ignore=False
# (CLI: --include-ignored) to scan everything regardless of these declarations.
_IGNORE_FILES = (".gitignore", ".forceignore")


def _load_ignore_specs(root: Path) -> dict[str, tuple[Path, object]]:
    """Load .gitignore / .forceignore as pathspec matchers, searching upward.

    SFDX ignore files conventionally live at the project root, but users often
    point graphify at a subdir (e.g. ``force-app``). Like git/sf, we walk up from
    ``root`` to the filesystem root to find the nearest ``.gitignore`` and
    ``.forceignore``. Each spec is returned with the directory it was found in, so
    path matching is anchored relative to that base (correct gitwildmatch semantics).

    Returns {filename: (base_dir, PathSpec)}.
    """
    specs: dict[str, tuple[Path, object]] = {}
    if pathspec is None:
        return specs
    for name in _IGNORE_FILES:
        d = root
        while True:
            f = d / name
            if f.exists():
                try:
                    lines = f.read_text(encoding="utf-8", errors="ignore").splitlines()
                    specs[name] = (d, pathspec.PathSpec.from_lines("gitignore", lines))
                except OSError:
                    pass
                break
            if d.parent == d:
                break
            d = d.parent
    return specs


def _matched_ignore_file(path: Path, root: Path, specs: dict[str, tuple[Path, object]]) -> str | None:
    """Return the name of the first ignore file whose patterns match `path`, else None."""
    if not specs:
        return None
    for name, (base, spec) in specs.items():
        try:
            rel = path.relative_to(base).as_posix()
        except ValueError:
            continue
        if spec.match_file(rel):  # type: ignore[attr-defined]
            return name
    return None


def extract_pdf_text(path: Path) -> str:
    """Extract plain text from a PDF using pypdf. Returns '' if pypdf not installed."""
    try:
        from pypdf import PdfReader

        reader = PdfReader(str(path))
        pages = [page.extract_text() or "" for page in reader.pages]
        return "\n".join(pages)
    except ImportError:
        return ""
    except Exception:
        return ""


def docx_to_markdown(path: Path) -> str:
    """Convert a .docx file to markdown text using python-docx. Returns '' if not installed."""
    try:
        from docx import Document

        doc = Document(str(path))
        lines = []
        for para in doc.paragraphs:
            style = para.style.name if para.style else ""
            text = para.text.strip()
            if not text:
                lines.append("")
                continue
            if style.startswith("Heading 1"):
                lines.append(f"# {text}")
            elif style.startswith("Heading 2"):
                lines.append(f"## {text}")
            elif style.startswith("Heading 3"):
                lines.append(f"### {text}")
            elif style.startswith("List"):
                lines.append(f"- {text}")
            else:
                lines.append(text)
        for table in doc.tables:
            rows = [[cell.text.strip() for cell in row.cells] for row in table.rows]
            if not rows:
                continue
            header = "| " + " | ".join(rows[0]) + " |"
            sep = "| " + " | ".join("---" for _ in rows[0]) + " |"
            lines.extend([header, sep])
            for row in rows[1:]:
                lines.append("| " + " | ".join(row) + " |")
        return "\n".join(lines)
    except ImportError:
        return ""
    except Exception:
        return ""


def xlsx_to_markdown(path: Path) -> str:
    """Convert an .xlsx file to markdown using openpyxl. Returns '' if not installed."""
    try:
        import openpyxl

        wb = openpyxl.load_workbook(str(path), read_only=True, data_only=True)
        sections = []
        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            rows = []
            for row in ws.iter_rows(values_only=True):
                if all(cell is None for cell in row):
                    continue
                rows.append([str(cell) if cell is not None else "" for cell in row])
            if not rows:
                continue
            sections.append(f"## Sheet: {sheet_name}")
            header = "| " + " | ".join(rows[0]) + " |"
            sep = "| " + " | ".join("---" for _ in rows[0]) + " |"
            sections.extend([header, sep])
            for row in rows[1:]:
                sections.append("| " + " | ".join(row) + " |")
        wb.close()
        return "\n".join(sections)
    except ImportError:
        return ""
    except Exception:
        return ""


def convert_office_file(path: Path, out_dir: Path) -> Path | None:
    """Convert .docx or .xlsx to a markdown sidecar in out_dir.
    Returns the sidecar path, or None if conversion failed/library not installed.
    """
    ext = path.suffix.lower()
    if ext == ".docx":
        text = docx_to_markdown(path)
    elif ext == ".xlsx":
        text = xlsx_to_markdown(path)
    else:
        return None
    if not text.strip():
        return None
    out_dir.mkdir(parents=True, exist_ok=True)
    name_hash = hashlib.sha256(str(path.resolve()).encode()).hexdigest()[:8]
    out_path = out_dir / f"{path.stem}_{name_hash}.md"
    out_path.write_text(f"<!-- converted from {path.name} -->\n\n{text}", encoding="utf-8")
    return out_path


def _classify_doc_file(path: Path) -> DocFileType | None:
    """Classify a file as a doc/paper/image type. Returns None if not a doc file."""
    ext = path.suffix.lower()
    if ext in _DOC_EXTENSIONS:
        return _DOC_EXTENSIONS[ext]
    if ext in _PAPER_EXTENSIONS:
        return _PAPER_EXTENSIONS[ext]
    if ext in _OFFICE_EXTENSIONS:
        return _OFFICE_EXTENSIONS[ext]
    if ext in _IMAGE_EXTENSIONS:
        return _IMAGE_EXTENSIONS[ext]
    return None


def detect(root: Path, *, respect_ignore: bool = True) -> dict:
    """Scan the SFDX project and return detected files grouped by type.

    When ``respect_ignore`` is True (default), files matching the project's
    ``.gitignore`` or ``.forceignore`` are skipped (gitwildmatch semantics).
    Pass ``respect_ignore=False`` (CLI: ``--include-ignored``) to scan everything
    regardless of those declarations. The hardcoded ``_SKIP_DIRS`` floor and the
    ``.graphifysfignore`` file are always applied.
    """
    root = root.resolve()
    files = {ft.value: [] for ft in SFFileType}
    doc_files: dict[str, list[str]] = {ft.value: [] for ft in DocFileType}
    bundle_dirs = {"lwc": [], "aura": []}
    skipped = []
    # Per-source skip counts for an honest summary line (never silently drop).
    skipped_by_source: dict[str, int] = {".gitignore": 0, ".forceignore": 0, ".graphifysfignore": 0}
    ignore_patterns = _load_sfgraphignore(root)
    ignore_specs = _load_ignore_specs(root) if respect_ignore else {}
    converted_dir = root / "graphify-sf-out" / "converted"

    def _skip_reason(path: Path) -> str | None:
        """Return the ignore-source name if path should be skipped, else None."""
        if _is_ignored(path, root, ignore_patterns):
            return ".graphifysfignore"
        return _matched_ignore_file(path, root, ignore_specs)

    for dirpath, dirnames, filenames in os.walk(root):
        dp = Path(dirpath)

        # Prune skip dirs
        dirnames[:] = [d for d in dirnames if d not in _SKIP_DIRS]

        # Check for LWC/Aura bundles
        if dp.parent.name == "lwc" and _is_lwc_bundle(dp):
            reason = _skip_reason(dp)
            if reason is None:
                bundle_dirs["lwc"].append(str(dp))
            else:
                skipped.append(str(dp))
                skipped_by_source[reason] = skipped_by_source.get(reason, 0) + 1
            continue

        if dp.parent.name == "aura" and _is_aura_bundle(dp):
            reason = _skip_reason(dp)
            if reason is None:
                bundle_dirs["aura"].append(str(dp))
            else:
                skipped.append(str(dp))
                skipped_by_source[reason] = skipped_by_source.get(reason, 0) + 1
            continue

        # Classify individual files
        for fname in filenames:
            path = dp / fname
            reason = _skip_reason(path)
            if reason is not None:
                skipped.append(str(path))
                skipped_by_source[reason] = skipped_by_source.get(reason, 0) + 1
                continue
            # Try SF classification first
            ftype = _classify_file(path)
            if ftype:
                files[ftype.value].append(str(path))
                continue
            # Try doc classification
            dtype = _classify_doc_file(path)
            if dtype:
                ext = path.suffix.lower()
                if ext in _OFFICE_EXTENSIONS:
                    # Convert office files to markdown sidecars
                    try:
                        md_path = convert_office_file(path, converted_dir)
                        if md_path:
                            doc_files[dtype.value].append(str(md_path))
                        else:
                            print(
                                f"[graphify-sf] WARNING: {path.name} skipped — "
                                "install graphify-sf[docs] to enable office file support",
                                file=sys.stderr,
                            )
                    except Exception as exc:
                        print(f"[graphify-sf] WARNING: office conversion failed for {path}: {exc}", file=sys.stderr)
                else:
                    doc_files[dtype.value].append(str(path))

    total_files = (
        sum(len(v) for v in files.values())
        + len(bundle_dirs["lwc"])
        + len(bundle_dirs["aura"])
        + sum(len(v) for v in doc_files.values())
    )

    return {
        "files": files,
        "doc_files": doc_files,
        "total_files": total_files,
        "bundle_dirs": bundle_dirs,
        "warning": None,
        "skipped": skipped,
        "skipped_count": len(skipped),
        "skipped_by_source": skipped_by_source,
        "respect_ignore": respect_ignore,
    }


def _md5_file(path: Path) -> str:
    """MD5 hash of file content for change detection."""
    h = hashlib.md5(usedforsecurity=False)
    try:
        with path.open("rb") as f:
            for chunk in iter(lambda: f.read(65536), b""):
                h.update(chunk)
    except OSError:
        return ""
    return h.hexdigest()


def load_manifest(manifest_path: str) -> dict:
    """Load the manifest from a previous run."""
    try:
        return json.loads(Path(manifest_path).read_text(encoding="utf-8"))
    except Exception:
        return {}


def save_manifest(
    files: dict[str, list[str]], manifest_path: str, doc_files: dict[str, list[str]] | None = None
) -> None:
    """Save current file mtimes + content hashes for incremental updates."""
    manifest = {}
    all_file_lists = list(files.values())
    if doc_files:
        all_file_lists.extend(doc_files.values())
    for file_list in all_file_lists:
        for f in file_list:
            try:
                p = Path(f)
                manifest[f] = {"mtime": p.stat().st_mtime, "hash": _md5_file(p)}
            except OSError:
                pass
    Path(manifest_path).parent.mkdir(parents=True, exist_ok=True)
    Path(manifest_path).write_text(json.dumps(manifest, indent=2), encoding="utf-8")


def detect_incremental(root: Path, manifest_path: str, *, respect_ignore: bool = True) -> dict:
    """Like detect(), but returns only new or modified files since the last run."""
    full = detect(root, respect_ignore=respect_ignore)
    manifest = load_manifest(manifest_path)

    if not manifest:
        full["new_files"] = full["files"]
        full["unchanged_files"] = {k: [] for k in full["files"]}
        full["new_doc_files"] = full["doc_files"]
        full["unchanged_doc_files"] = {k: [] for k in full["doc_files"]}
        return full

    new_files = {k: [] for k in full["files"]}
    unchanged_files = {k: [] for k in full["files"]}
    new_doc_files = {k: [] for k in full["doc_files"]}
    unchanged_doc_files = {k: [] for k in full["doc_files"]}

    def _check_changed(f: str) -> bool:
        stored = manifest.get(f, {})
        try:
            current_mtime = Path(f).stat().st_mtime
            stored_mtime = stored.get("mtime")
            if stored_mtime is None or current_mtime != stored_mtime:
                return _md5_file(Path(f)) != stored.get("hash", "")
            return False
        except Exception:
            return True

    for ftype, file_list in full["files"].items():
        for f in file_list:
            if _check_changed(f):
                new_files[ftype].append(f)
            else:
                unchanged_files[ftype].append(f)

    for dtype, file_list in full["doc_files"].items():
        for f in file_list:
            if _check_changed(f):
                new_doc_files[dtype].append(f)
            else:
                unchanged_doc_files[dtype].append(f)

    full["new_files"] = new_files
    full["unchanged_files"] = unchanged_files
    full["new_doc_files"] = new_doc_files
    full["unchanged_doc_files"] = unchanged_doc_files
    return full
